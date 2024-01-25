# This script takes the exported unlinked concepts and groups them by
# category (Noun, Adjective, etc). It also groups proper nouns separately.
# The concepts are sorted alphabetically and put into tables in a Word
# document which can be sent to the MTT to translate.
#
# If no concepts for a particular category are found, that table is excluded.
#
# Example sentences are drawn from the verse that includes the concept. If
# no sentence is found to contain that concept, the whole verse is shown
# and highlighted so that the user can attend to it before sending to the MTT.

import sys
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.shared import Pt, Cm


class UnlinkedConcept:
    def __init__(self, group, word, gloss):
        self.group = group
        self.word = word
        self.gloss = gloss
        self.verse = ''
        self.sample = ''


def get_params():
    if len(sys.argv) < 2:
        print('Please specify a .txt file to import')
        return None

    file_name = sys.argv[-1]
    if file_name.startswith('-'):
        print('File name must be the last argument')
        return None

    file_path = Path(file_name).with_suffix('.txt')
    if not file_path.exists():
        print('Specified File does not exist...')
        return None

    return {
        'input_path': file_path,
        'output_path': file_path.with_name(f'Lexicon - {file_path.stem}.docx'),
        'add_notes_column': '-N' in sys.argv or '-n' in sys.argv
    }


def import_concepts(params):
    CONCEPT_REGEX = re.compile(r'^Concept \(([a-zA-Z]+)\): ([.a-zA-Z0-9- ]+?-[A-Z])(?:  \'(.+?)\')?$')
    groups = {}

    path = params['input_path']
    with path.open() as f:
        concept = None
        for line_num, line in enumerate(f):
            if line.startswith('Concept'):
                m = CONCEPT_REGEX.match(line)
                if not m:
                    print('Unexpected format for Concept on line ' + str(line_num))
                    continue
                word = m[2]
                group  = 'Proper Nouns' if m[1] == 'Noun' and word[0].isupper() else (m[1] + 's')
                gloss = m[3] or ''
                concept = UnlinkedConcept(group, word, gloss)

            elif line.startswith('Sample Sentence'):
                concept.sample = line[len('Sample Sentence: '):].strip()

            elif line.startswith('Verse'):
                concept.verse = line[len('Verse: '):].strip()
                if concept.group not in groups:
                    groups[concept.group] = []
                groups[concept.group].append(concept)

            elif line.startswith('Current Passage'):
                params['passage'] = line[len('Current Passage: '):].strip()

    for k,v in groups.items():
        print(f'Retrieved {len(v)} {k} unlinked concepts from "{path}"')
    return groups


def export_document(groups, params):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri (Body)'

    # Set orientation to landscape
    section = doc.sections[-1]
    old_width, old_height = section.page_width, section.page_height
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = old_width
    section.page_width = old_height

    # Set the margins to Normal (2.54cm for each)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Add the passage
    p = doc.add_paragraph(params['passage'])
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(14)

    # Create the tables
    add_notes_column = params['add_notes_column']
    next_idx = create_names_tables(groups, doc)
    next_idx = create_table('Nouns', groups, doc, next_idx, add_notes_column)
    next_idx = create_table('Verbs', groups, doc, next_idx, add_notes_column)
    next_idx = create_table('Adjectives', groups, doc, next_idx, add_notes_column)
    next_idx = create_table('Adverbs', groups, doc, next_idx, add_notes_column)
    next_idx = create_table('Adpositions', groups, doc, next_idx, add_notes_column)
    next_idx = create_table('Conjunctions', groups, doc, next_idx, add_notes_column)
    next_idx = create_table('Particles', groups, doc, next_idx, add_notes_column)
    doc.save(str(params['output_path']))


def create_names_tables(groups, doc):
    key = 'Proper Nouns'
    if key not in groups:
        return 1

    # This table holds all the Proper Nouns
    add_table_caption("Proper Names", doc, 1)
    table = doc.add_table(rows=1, cols=3, style='Table Grid')
    add_table_headers("Proper Names", table, proper=True)
    set_column_widths(table, [Cm(5)] * 3)

    for concept in groups[key]:
        row_cells = table.add_row().cells
        row_cells[0].text = concept.word
        row_cells[1].text = concept.gloss
    set_column_widths(table, [Cm(5)] * 3)

    return 2


def create_table(group_name, groups, doc, idx, add_notes_column):
    if group_name not in groups:
        return idx
    add_table_caption(group_name, doc, idx)
    table = doc.add_table(rows=1, cols=6, style='Table Grid')
    add_table_headers(group_name, table)
    
    for concept in sorted(groups[group_name], key=lambda c: c.word):
        row_cells = table.add_row().cells
        row_cells[0].text = concept.word
        row_cells[1].text = concept.gloss
        add_verse_sentences(concept, row_cells[2])
        add_sample_sentences(concept, row_cells[3])

    col_widths = []
    if add_notes_column:
        col_widths = [Cm(2.5), Cm(3), Cm(6), Cm(3), Cm(3), Cm(3), Cm(3)]
        table.add_column(col_widths[-1])
        table.cell(0, len(col_widths)-1).text = 'Notes'
    else:
        col_widths = [Cm(2.7), Cm(3.2), Cm(6.3), Cm(4), Cm(3.4), Cm(3.4)]

    set_column_widths(table, col_widths)
    return idx + 1


def add_table_caption(name, doc, idx):
    if idx > 1:
        doc.add_paragraph().paragraph_format.space_after = Pt(0)
    caption = doc.add_paragraph(f'Table {idx}. {name}')
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(0)


def add_table_headers(name, table, proper=False):
    names = []
    if proper:
        names = ['Nouns: ' + name, 'Glosses', 'Target Words']
    else:
        names = [
            name,
            'Glosses',
            'Verses',
            'Target Sentences',
            'Target Words',
            'Target Glosses'
        ]

    header_cols = table.row_cells(0)
    if len(names) != len(header_cols):
        # This should never happen but safer to check anyway
        print('Wrong number of columns for ' + name)
        return

    for idx, text in enumerate(names):
        if text == 'Target Words':
            header_cols[idx].text = text
            header_cols[idx].paragraphs[0].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            header_cols[idx].text = text


SENTENCE_REGEX = re.compile(r'([^.?!]+[.?!]\S*) ?')
def add_verse_sentences(concept, table_cell):
    # Remove the sense -X from the word
    dash_idx = concept.word.rindex('-')
    word = concept.word[:dash_idx]
    
    # Split the reference and verse text
    colon_idx = concept.verse.index(':')
    verse_start = concept.verse.index(' ', colon_idx)
    ref = concept.verse[:verse_start]
    verse_text = concept.verse[verse_start+1:]

    # Start with the verse ref
    table_cell.text = ref
    paragraph = table_cell.paragraphs[0]

    # Split the sentences and return the ones that contain the word
    word_regex = r'\b(' + re.escape(word) + r'(?:s|es|ed|d)?)\b'
    word_found = False
    for sentence in SENTENCE_REGEX.findall(verse_text):
        m = re.search(word_regex, sentence, re.IGNORECASE)
        if m:
            # Make the matched word bold in the sentence
            start, end = m.span(1)
            paragraph.add_run(text=' ' + sentence[:start])
            paragraph.add_run(text=sentence[start:end]).bold = True
            paragraph.add_run(text=sentence[end:])
            word_found = True

    if not word_found:
        # The verse might have an unrecognizable form of the word, the word
        # might not be present at all due to restructuring. If the word is
        # not found, show the whole verse and highlight the text so the user
        # knows to attend to it
        paragraph.clear()
        run_font = paragraph.add_run(text=concept.verse).font
        run_font.highlight_color = WD_COLOR_INDEX.YELLOW
        run_font.size = Pt(10)
    else:
        for run in paragraph.runs:
            run.font.size = Pt(10)


def add_sample_sentences(concept, table_cell):
    if not concept.sample:
        return

    # Start with the sample sentence itself with the separating bar
    table_cell.text = concept.sample + ' | '
    paragraph = table_cell.paragraphs[0]

    # Add the place where the translator needs to translate
    run_font = paragraph.add_run(text='Translation here.').font
    run_font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Set the font size
    for run in paragraph.runs:
        run.font.size = Pt(10)


def set_column_widths(table, widths):
    for idx, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = widths[idx]


if __name__ == "__main__":
    params = get_params()
    if params:
        verses = import_concepts(params)
        export_document(verses, params)
        print(f'Successfully exported "{params["output_path"]}"')
